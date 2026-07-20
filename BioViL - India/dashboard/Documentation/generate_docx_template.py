"""
BioViL Corporate Identity — DOCX Template Generator
Generates a reusable Word document template with branded styles,
headers, footers, and placeholder sections for future reports.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = "/Users/pankajsharma/Library/CloudStorage/OneDrive-Personal/AI Agents/Bioenergy_KPI_Dashboard/Logo/logo 1colour  bg w.svg"

# ─── BioViL CI Colors ───
TEAL = RGBColor(0x08, 0x3F, 0x5E)
ORANGE = RGBColor(0xF7, 0x94, 0x1D)
EMERALD = RGBColor(0x10, 0xB9, 0x81)
SLATE = RGBColor(0x33, 0x41, 0x55)
LIGHT_GRAY = RGBColor(0x94, 0xA3, 0xB8)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_shading(cell, color_hex):
    """Apply background color to a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_bottom_border(paragraph, color="083F5E", size="12"):
    """Add a bottom border to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def create_template():
    doc = Document()
    
    # ─── PAGE SETUP ───
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    
    # ─── DEFINE STYLES ───
    style = doc.styles
    
    # Title Style
    title_style = style['Title']
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(28)
    title_font.bold = True
    title_font.color.rgb = TEAL
    title_style.paragraph_format.space_after = Pt(6)
    
    # Heading 1
    h1 = style['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = TEAL
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(8)
    
    # Heading 2
    h2 = style['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = TEAL
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    
    # Heading 3
    h3 = style['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = SLATE
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)
    
    # Normal body
    normal = style['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.font.color.rgb = SLATE
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.5
    
    # ─── HEADER ───
    header = section.header
    header.is_linked_to_previous = False
    htable = header.add_table(1, 2, width=Inches(6.5))
    htable.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Left cell: Company name
    left_cell = htable.cell(0, 0)
    left_cell.width = Inches(3.5)
    lp = left_cell.paragraphs[0]
    lr = lp.add_run('BioViL Energy')
    lr.font.name = 'Calibri'
    lr.font.size = Pt(10)
    lr.font.bold = True
    lr.font.color.rgb = TEAL
    
    # Right cell: Document type placeholder
    right_cell = htable.cell(0, 1)
    right_cell.width = Inches(3.0)
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run('[DOCUMENT TYPE]')
    rr.font.name = 'Calibri'
    rr.font.size = Pt(9)
    rr.font.color.rgb = ORANGE
    rr.font.bold = True
    
    # Add header border line
    hp = header.add_paragraph()
    add_bottom_border(hp, "083F5E", "12")
    
    # ─── FOOTER ───
    footer = section.footer
    footer.is_linked_to_previous = False
    ftable = footer.add_table(1, 2, width=Inches(6.5))
    ftable.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    fl = ftable.cell(0, 0).paragraphs[0]
    flr = fl.add_run('BioViL Energy — Confidential')
    flr.font.name = 'Calibri'
    flr.font.size = Pt(8)
    flr.font.color.rgb = LIGHT_GRAY
    
    fr = ftable.cell(0, 1).paragraphs[0]
    fr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    frr = fr.add_run('Page ')
    frr.font.name = 'Calibri'
    frr.font.size = Pt(8)
    frr.font.color.rgb = LIGHT_GRAY
    # Add auto page number field (Properly wrapped in a Run to prevent corruption)
    page_run = fr.add_run()
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    page_run._r.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    page_run._r.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    page_run._r.append(fldChar2)
    
    # ─── COVER PAGE CONTENT ───
    # Add some spacing
    for _ in range(6):
        doc.add_paragraph()
    
    # Section tag
    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_run = tag.add_run('TECHNICAL REPORT')
    tag_run.font.name = 'Calibri'
    tag_run.font.size = Pt(10)
    tag_run.font.color.rgb = ORANGE
    tag_run.font.bold = True
    tag_run.font.letter_spacing = Pt(3)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('[Report Title]')
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = TEAL
    
    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run('[Subtitle or description of the report]')
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = LIGHT_GRAY
    
    # Date & Confidentiality
    for _ in range(4):
        doc.add_paragraph()
    
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_bottom_border(meta, "E2E8F0", "4")
    meta_run = meta.add_run('[Month Year]  •  Confidential  •  Version [X.X]')
    meta_run.font.name = 'Calibri'
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = LIGHT_GRAY
    
    # Page break after cover
    doc.add_page_break()
    
    # ─── SAMPLE BODY SECTIONS ───
    
    # Section 1
    sec_tag = doc.add_paragraph()
    sec_run = sec_tag.add_run('SECTION 01')
    sec_run.font.name = 'Calibri'
    sec_run.font.size = Pt(10)
    sec_run.font.color.rgb = ORANGE
    sec_run.font.bold = True
    
    doc.add_heading('[Section Title]', level=1)
    doc.add_paragraph('[Introductory paragraph text goes here. Replace this placeholder with your content. The font, sizing, and spacing are pre-configured to match the BioViL Corporate Identity.]')
    
    # Sample KPI Table
    doc.add_heading('Key Metrics', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    hdr_cells = table.rows[0].cells
    headers = ['Metric', 'Value', 'Unit', 'Status']
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        set_cell_shading(hdr_cells[i], '083F5E')
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = WHITE
                r.font.bold = True
                r.font.size = Pt(9)
                r.font.name = 'Calibri'
    
    # Sample data rows
    sample_data = [
        ['CO₂e Avoided', '438.1', 'kg/month', '✓ On Track'],
        ['Waste Diverted', '600', 'kg/month', '✓ On Track'],
        ['Biogas Produced', '~20', 'm³/month', '✓ On Track'],
    ]
    for row_data in sample_data:
        row = table.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = text
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = 'Calibri'
                    r.font.color.rgb = SLATE
    
    # Callout box (simulated with a single-cell table)
    doc.add_paragraph()
    callout_label = doc.add_paragraph()
    callout_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cl_run = callout_label.add_run('▎ KEY INSIGHT')
    cl_run.font.name = 'Calibri'
    cl_run.font.size = Pt(9)
    cl_run.font.color.rgb = EMERALD
    cl_run.font.bold = True
    
    callout = doc.add_paragraph()
    callout_run = callout.add_run('[Place important callouts, formulae, or key data points here. This styled block draws attention to critical information.]')
    callout_run.font.name = 'Calibri'
    callout_run.font.size = Pt(10)
    callout_run.font.color.rgb = SLATE
    callout_run.font.italic = True
    
    # Another section
    doc.add_page_break()
    
    sec_tag2 = doc.add_paragraph()
    sec_run2 = sec_tag2.add_run('SECTION 02')
    sec_run2.font.name = 'Calibri'
    sec_run2.font.size = Pt(10)
    sec_run2.font.color.rgb = ORANGE
    sec_run2.font.bold = True
    
    doc.add_heading('[Next Section Title]', level=1)
    doc.add_paragraph('[Continue your report content here. Insert figures, tables, and charts as needed. All styling will inherit the BioViL CI automatically.]')
    
    doc.add_heading('[Subsection]', level=2)
    doc.add_paragraph('[Subsection content with supporting data and analysis.]')
    
    doc.add_heading('[Sub-subsection]', level=3)
    doc.add_paragraph('[Detailed technical explanation or methodology notes.]')
    
    # References section
    doc.add_page_break()
    
    sec_tag3 = doc.add_paragraph()
    sec_run3 = sec_tag3.add_run('REFERENCES')
    sec_run3.font.name = 'Calibri'
    sec_run3.font.size = Pt(10)
    sec_run3.font.color.rgb = ORANGE
    sec_run3.font.bold = True
    
    doc.add_heading('Source Documents & Literature', level=1)
    
    refs = [
        '[1]  Author, A. (Year). Title of document. Publisher/Source.',
        '[2]  Author, B. (Year). Title of second reference. Journal Name, Volume(Issue), Pages.',
        '[3]  Organization. (Year). Title of report. URL or DOI.',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        r = p.add_run(ref)
        r.font.name = 'Calibri'
        r.font.size = Pt(9.5)
        r.font.color.rgb = LIGHT_GRAY
    
    # ─── SAVE ───
    output_path = os.path.join(DOCS_DIR, 'BioViL_Report_Template.docx')
    doc.save(output_path)
    print(f"✓ DOCX Template saved: {output_path}")


if __name__ == "__main__":
    create_template()
